#!/usr/bin/env python3
"""Streamlit UI for the deployed TY planning tool."""

from __future__ import annotations

import csv
from datetime import datetime
import hashlib
import html
from io import BytesIO
import json
import logging
import os
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path
from urllib import error as urllib_error
from urllib import request as urllib_request
import zipfile

import streamlit as st


ROOT = Path(__file__).resolve().parents[1]
SCRIPTS_DIR = ROOT / "scripts"
GENERATED_PLANS_DIR = ROOT / "outputs" / "generated_plans"
LEADS_DIR = ROOT / "outputs" / "leads"
SHOW_DOCX_DIAGNOSTIC = False
SHOW_PDF_DIAGNOSTIC = False
SHOW_PREVIEW_DIAGNOSTIC = False
LOGGER = logging.getLogger(__name__)
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

from answer_query import (  # noqa: E402
    answer_question,
    heading_aliases,
    normalise_school_display_name,
    normalise_template_context,
    parse_template_context,
    requested_output_language,
)


def split_layer_prefix(text: str) -> tuple[str | None, str]:
    match = re.match(r"^\[(official_guidance|ciunas_framework)\]\s+(.*)$", text.strip())
    if not match:
        return None, text.strip()
    return match.group(1), match.group(2).strip()


def render_source_layer_caption() -> None:
    st.caption(
        "Official guidance informs policy and planning context. "
        "Ciunas framework sources can shape programme and module ideas, but they are not official policy."
    )


def plan_sections_map(full_plan_text: str) -> tuple[str, str, dict[str, list[str]]]:
    title, subtitle, sections = parse_plan_blocks(full_plan_text)
    section_map: dict[str, list[str]] = {}
    for heading, paragraphs in sections:
        if heading.strip():
            section_map[heading.strip().lower()] = paragraphs
    return title, subtitle, section_map


def extract_preview_payload(answer: str) -> tuple[str, str, str, list[str]]:
    title, subtitle, _sections = parse_plan_blocks(answer)
    language = infer_plan_language(title, subtitle)
    canonical_headings = (
        ["Programme Overview", "Rationale", "Aims"]
        if language == "en"
        else ["Forbhreathnú ar an gClár", "Réasúnaíocht", "Aidhmeanna"]
    )
    alias_map = heading_aliases(language)
    all_heading_variants = set()
    for canonical_heading, aliases in alias_map.items():
        all_heading_variants.add(canonical_heading.strip().lower())
        all_heading_variants.update(alias.lower() for alias in aliases)

    lines = [line.rstrip() for line in answer.splitlines()]
    non_empty_lines = [line.strip() for line in lines if line.strip()]
    title_line = non_empty_lines[0] if non_empty_lines else ""
    subtitle_line = non_empty_lines[1] if len(non_empty_lines) > 1 else ""

    line_items: list[tuple[str, str]] = []
    for line in lines:
        stripped = line.strip()
        if stripped:
            line_items.append((line, stripped))

    start_index = 0
    if title_line:
        for idx, (_line, stripped) in enumerate(line_items):
            if stripped == title_line:
                start_index = idx + 1
                break
    if subtitle_line:
        for idx in range(start_index, len(line_items)):
            if line_items[idx][1] == subtitle_line:
                start_index = idx + 1
                break

    candidate_lines = line_items[start_index:]
    intro_lines: list[str] = []
    for _raw_line, stripped in candidate_lines:
        if stripped.lower() in all_heading_variants:
            break
        intro_lines.append(stripped)
    matched_sections: list[tuple[str, list[str]]] = []
    matched_names: list[str] = []

    for canonical_heading in canonical_headings:
        variants = {canonical_heading.strip().lower(), *[alias.lower() for alias in alias_map.get(canonical_heading, [])]}
        match_index = None
        match_heading = canonical_heading
        for idx, (_raw_line, stripped) in enumerate(candidate_lines):
            if stripped.lower() in variants:
                match_index = idx
                match_heading = stripped
                break
        if match_index is None:
            continue

        body_lines: list[str] = []
        remainder = candidate_lines[match_index + 1 :]
        for _raw_line, stripped in remainder:
            if stripped.lower() in all_heading_variants:
                break
            body_lines.append(stripped)

        body_text = " ".join(body_lines).strip()
        matched_sections.append((match_heading, [body_text] if body_text else []))
        matched_names.append(canonical_heading)
        candidate_lines = remainder

    if not matched_sections:
        preview_sections = parse_plan_blocks(answer)[2][:3]
    else:
        preview_sections = matched_sections

    preview_text_parts: list[str] = []
    if title:
        preview_text_parts.append(f"### {title}")
    if subtitle:
        preview_text_parts.append(subtitle)
    if intro_lines:
        preview_text_parts.append("")
        preview_text_parts.append(" ".join(intro_lines).strip())
    for heading, paragraphs in preview_sections:
        preview_text_parts.append("")
        preview_text_parts.append(f"#### {heading}")
        preview_text_parts.extend(paragraphs)
    return title, subtitle, "\n\n".join(part for part in preview_text_parts if part is not None).strip(), matched_names


def describe_pdf_runtime_support() -> tuple[bool, str]:
    try:
        from reportlab.lib.pagesizes import A4  # noqa: F401
        from reportlab.platypus import SimpleDocTemplate  # noqa: F401

        return True, "ReportLab backend available."
    except Exception as exc:
        lualatex_path = shutil.which("lualatex")
        if lualatex_path:
            return True, f"LuaLaTeX backend available at {lualatex_path}."
        return False, f"ReportLab import failed: {exc.__class__.__name__}: {exc}"


def render_plan_sections(title: str, subtitle: str, sections: list[tuple[str, list[str]]]) -> None:
    if not title and not sections:
        st.write("No TY plan text was returned.")
        return

    if title:
        st.markdown(f"### {title}")
    if subtitle:
        st.caption(subtitle)

    for heading, paragraphs in sections:
        if heading:
            st.markdown(f"#### {heading}")
        for paragraph in paragraphs:
            if paragraph:
                st.write(paragraph)


def render_generated_plan(answer: str) -> None:
    title, subtitle, sections = parse_plan_blocks(answer)
    render_plan_sections(title, subtitle, sections)


def render_plan_preview(answer: str) -> None:
    title, subtitle, preview_text, matched_names = extract_preview_payload(answer)
    if preview_text:
        st.markdown(preview_text)
    else:
        render_plan_sections(title, subtitle, [])
    if SHOW_PREVIEW_DIAGNOSTIC:
        st.caption(f"Temporary diagnostic: Preview length {len(preview_text)} characters.")
        st.caption(
            "Temporary diagnostic: Preview sections detected: "
            + (", ".join(matched_names) if matched_names else "none")
        )


def append_coordinator_name(prompt: str, coordinator_name: str = "") -> str:
    cleaned_name = coordinator_name.strip()
    if not cleaned_name:
        return prompt
    return f"{prompt.rstrip()}\nTY Coordinator: {cleaned_name}"


def build_download_prompt(user_input: str, coordinator_name: str = "") -> str:
    if not user_input:
        return append_coordinator_name("Create a TY plan", coordinator_name)

    lowered = user_input.lower()
    if any(term in lowered for term in ("create", "generate", "plan", "ty plan", "annual plan")):
        return append_coordinator_name(user_input, coordinator_name)
    if any(term in lowered for term in ("irish", "gaeilge", "as gaeilge", "i ngaeilge", "cruthaigh", "idirbhliana")):
        return append_coordinator_name(f"Create a TY plan in Irish focused on {user_input}", coordinator_name)
    return append_coordinator_name(f"Create a TY plan focused on {user_input}", coordinator_name)


def build_tailored_plan_prompt(
    school_name: str,
    cohort_size: str,
    coordinator_name: str,
    school_type: str,
    school_ethos: str,
    priorities: str,
    existing_modules: str,
    work_experience: str,
    additional_context: str,
    language: str,
) -> str:
    return (
        "Create a Transition Year annual plan.\n\n"
        "Use the following context:\n"
        f"School name: {school_name or 'Not specified'}\n"
        f"Cohort size: {cohort_size or 'Not specified'}\n"
        f"TY Coordinator: {coordinator_name or 'Not specified'}\n"
        f"School type: {school_type or 'Not specified'}\n"
        f"School ethos: {school_ethos or 'Not specified'}\n"
        f"Main priorities: {priorities or 'Not specified'}\n"
        f"Existing modules: {existing_modules or 'Not specified'}\n"
        f"Work experience timing: {work_experience or 'Not specified'}\n"
        f"Additional context: {additional_context or 'Not specified'}\n"
        f"Language: {language or 'English'}\n\n"
        "The plan should:\n"
        "- reflect the school context\n"
        "- use the school name and cohort context naturally where relevant\n"
        "- prioritise the stated focus areas\n"
        "- incorporate existing modules where relevant\n"
        "- structure work experience appropriately\n"
        "- sound realistic for current TY practice in Irish schools\n"
        "- remain practical and usable for a TY coordinator\n\n"
        "Generate a full structured TY annual plan."
    )


def infer_output_language(answer_mode: str, prompt: str) -> str:
    context = parse_template_context(prompt)
    explicit_language = requested_output_language(context.get("language", ""))
    if explicit_language:
        return explicit_language
    lowered_prompt = prompt.lower()
    if answer_mode.endswith("_ga") or any(
        term in lowered_prompt for term in ("irish", "gaeilge", "as gaeilge", "i ngaeilge")
    ):
        return "ga"
    return "en"


def save_generated_plan(answer: str, language: str) -> tuple[Path, Path]:
    GENERATED_PLANS_DIR.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    stem = f"ty_annual_plan_{language}_{timestamp}"
    markdown_path = GENERATED_PLANS_DIR / f"{stem}.md"
    text_path = GENERATED_PLANS_DIR / f"{stem}.txt"
    markdown_path.write_text(answer + "\n", encoding="utf-8")
    text_path.write_text(answer + "\n", encoding="utf-8")
    return markdown_path, text_path


def save_export_file(data: bytes, language: str, format_marker: str, suffix: str) -> Path:
    GENERATED_PLANS_DIR.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    export_path = GENERATED_PLANS_DIR / f"ty_annual_plan_{language}_{format_marker}_{timestamp}.{suffix}"
    export_path.write_bytes(data)
    return export_path


def looks_like_email(value: str) -> bool:
    return bool(re.fullmatch(r"[^@\s]+@[^@\s]+\.[^@\s]+", value.strip()))


def save_lead_row(row: dict[str, str]) -> Path:
    LEADS_DIR.mkdir(parents=True, exist_ok=True)
    leads_path = LEADS_DIR / "ty_planner_leads.csv"
    fieldnames = [
        "timestamp",
        "email",
        "name",
        "mode",
        "language",
        "school_name",
        "school_type",
        "priorities",
    ]
    write_header = not leads_path.exists()
    with leads_path.open("a", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        if write_header:
            writer.writeheader()
        writer.writerow({field: row.get(field, "") for field in fieldnames})
    return leads_path


def read_secret_or_env(name: str) -> str:
    value = os.getenv(name, "").strip()
    if value:
        return value

    try:
        secret_value = st.secrets.get(name, "")
    except Exception:
        secret_value = ""

    return str(secret_value).strip()


def sync_mailchimp_lead(email: str, name: str | None = None) -> tuple[bool, str]:
    api_key = read_secret_or_env("MAILCHIMP_API_KEY")
    audience_id = read_secret_or_env("MAILCHIMP_AUDIENCE_ID")
    if not api_key:
        return False, "MAILCHIMP_API_KEY is missing."
    if not audience_id:
        return False, "MAILCHIMP_AUDIENCE_ID is missing."
    if "-" not in api_key:
        return False, "MAILCHIMP_API_KEY is missing its datacenter suffix."

    datacenter = api_key.rsplit("-", 1)[-1].strip()
    if not datacenter:
        return False, "Could not derive Mailchimp datacenter from MAILCHIMP_API_KEY."

    cleaned_email = email.strip().lower()
    subscriber_hash = hashlib.md5(cleaned_email.encode("utf-8")).hexdigest()
    first_name = (name or "").strip().split()[0] if (name or "").strip() else ""
    base_url = f"https://{datacenter}.api.mailchimp.com/3.0"
    member_url = f"{base_url}/lists/{audience_id}/members/{subscriber_hash}"
    tags_url = f"{member_url}/tags"
    auth_header = f"apikey {api_key}"

    member_payload = {
        "email_address": cleaned_email,
        "status_if_new": "subscribed",
        "status": "subscribed",
        "merge_fields": {"FNAME": first_name},
    }
    tags_payload = {"tags": [{"name": "TY Planner Lead", "status": "active"}]}

    def send_json(url: str, method: str, payload: dict[str, object]) -> tuple[bool, str]:
        body = json.dumps(payload).encode("utf-8")
        req = urllib_request.Request(
            url,
            data=body,
            method=method,
            headers={
                "Authorization": auth_header,
                "Content-Type": "application/json",
            },
        )
        try:
            with urllib_request.urlopen(req, timeout=20) as response:
                status = getattr(response, "status", 200)
                if 200 <= status < 300:
                    return True, f"{method} {url} returned {status}."
                return False, f"{method} {url} returned unexpected status {status}."
        except urllib_error.HTTPError as exc:
            try:
                detail = exc.read().decode("utf-8", "replace")
            except Exception:
                detail = ""
            return False, f"{method} {url} failed with HTTP {exc.code}: {detail or exc.reason}"
        except Exception as exc:
            return False, f"{method} {url} failed: {exc.__class__.__name__}: {exc}"

    member_ok, member_detail = send_json(member_url, "PUT", member_payload)
    if not member_ok:
        return False, member_detail

    tags_ok, tags_detail = send_json(tags_url, "POST", tags_payload)
    if not tags_ok:
        return False, tags_detail

    return True, "Mailchimp lead synced and tagged."


def is_generator_result(answer_mode: str) -> bool:
    return answer_mode.startswith("template_generation")


def latex_escape(text: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    return "".join(replacements.get(char, char) for char in text)


def split_export_paragraph(paragraph: str) -> list[str]:
    cleaned = " ".join(paragraph.split())
    if len(cleaned) < 520:
        return [cleaned]

    sentence_parts = re.split(r"(?<=[.!?])\s+", cleaned)
    sentence_parts = [part.strip() for part in sentence_parts if part.strip()]
    if len(sentence_parts) < 3:
        return [cleaned]

    midpoint = max(1, len(sentence_parts) // 2)
    left = " ".join(sentence_parts[:midpoint]).strip()
    right = " ".join(sentence_parts[midpoint:]).strip()
    if not left or not right:
        return [cleaned]
    return [left, right]


def infer_plan_language(title: str, subtitle: str) -> str:
    combined = f"{title} {subtitle}".lower()
    ga_markers = (
        "idirbhliana",
        "meán fómhair",
        "bealtaine",
        "plean bliantúil",
    )
    return "ga" if any(marker in combined for marker in ga_markers) else "en"


def standardised_export_subtitle(language: str) -> str:
    if language == "ga":
        return "Meán Fómhair 2026 – Bealtaine 2027"
    return "September 2026 – May 2027"


def build_cover_page_text(language: str) -> tuple[str, str, str, str]:
    if language == "ga":
        return (
            "PLEAN BLIANTÚIL",
            "NA hIDIRBHLIANA",
            "Comhordaitheoir TY:",
            "Doiciméad pleanála struchtúrtha do dhearadh, cur i bhfeidhm, agus athbhreithniú chlár na hIdirbhliana.",
        )
    return (
        "TRANSITION YEAR",
        "ANNUAL PLAN",
        "TY Coordinator:",
        "A structured planning document for the design, delivery, and review of the Transition Year programme.",
    )


def build_title_block_values(language: str, context: dict[str, str] | None = None) -> tuple[str, str, bool]:
    _, _, coordinator_label, _ = build_cover_page_text(language)
    context = normalise_template_context(context or {}, language)
    school_name = normalise_school_display_name(str(context.get("school_name", "")).strip())
    ty_coordinator = str(context.get("ty_coordinator", "")).strip()

    school_value = school_name if school_name else "______________________"
    coordinator_value = (
        f"{coordinator_label} {ty_coordinator}"
        if ty_coordinator
        else f"{coordinator_label} ______________________"
    )
    return school_value, coordinator_value, bool(school_name)


def needs_fill_lines(heading: str, language: str) -> bool:
    normalised = heading.strip().lower()
    if language == "ga":
        return normalised in {"croí-mhodúil", "féilire achomair", "punann"}
    return normalised in {"core modules", "summary calendar", "portfolio"}


def export_fill_lines() -> list[str]:
    return [
        "____________________________",
        "____________________________",
    ]


def parse_plan_blocks(full_plan_text: str) -> tuple[str, str, list[tuple[str, list[str]]]]:
    lines = [line.rstrip() for line in full_plan_text.splitlines()]
    non_empty_lines = [line.strip() for line in lines if line.strip()]
    if not non_empty_lines:
        return "", "", []

    title = non_empty_lines[0]
    subtitle = non_empty_lines[1] if len(non_empty_lines) > 1 else ""
    remaining_text = full_plan_text
    title_block = f"{title}\n{subtitle}".strip()
    if title_block and remaining_text.startswith(title_block):
        remaining_text = remaining_text[len(title_block) :].lstrip()

    blocks = [block.strip() for block in remaining_text.split("\n\n") if block.strip()]
    sections: list[tuple[str, list[str]]] = []

    for block in blocks:
        lines = [line.strip() for line in block.splitlines() if line.strip()]
        if len(lines) >= 2:
            heading = lines[0]
            body: list[str] = []
            for line in lines[1:]:
                body.extend(split_export_paragraph(line))
        else:
            heading = ""
            body = split_export_paragraph(block)
        sections.append((heading, body))

    return title, subtitle, sections


def clean_markdown_text(text: str) -> str:
    cleaned = text.strip()
    cleaned = re.sub(r"^\s*#{1,6}\s*", "", cleaned)
    cleaned = re.sub(r"^\s*\d+\.\s*", "", cleaned)
    cleaned = re.sub(r"^\s*[-*•]\s+", "", cleaned)
    cleaned = cleaned.replace("\\|", "|")
    cleaned = re.sub(r"`([^`]+)`", r"\1", cleaned)
    cleaned = re.sub(r"\*\*([^*]+)\*\*", r"\1", cleaned)
    cleaned = re.sub(r"__([^_]+)__", r"\1", cleaned)
    cleaned = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"\1", cleaned)
    cleaned = re.sub(r"(?<!_)_([^_]+)_(?!_)", r"\1", cleaned)
    cleaned = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned.strip()


def is_list_line(line: str) -> bool:
    return bool(re.match(r"^\s*[-*•]\s+", line))


def is_markdown_heading_line(line: str) -> bool:
    return bool(re.match(r"^\s*#{1,6}\s+\S+", line))


def is_numbered_heading_line(line: str) -> bool:
    stripped = line.strip()
    return bool(re.match(r"^\d+\.\s+\S+", stripped)) and len(stripped.split()) <= 12


def is_plain_heading_candidate(line: str) -> bool:
    stripped = clean_markdown_text(line)
    if not stripped:
        return False
    if any(token in stripped for token in ("|", ".", "?", "!")):
        return False
    if len(stripped) > 80:
        return False
    if len(stripped.split()) > 10:
        return False
    return True


def is_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.count("|") >= 2


def parse_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [clean_markdown_text(cell) for cell in stripped.split("|")]


def is_markdown_table_separator(row: list[str]) -> bool:
    if not row:
        return False
    compact = "".join(row).replace(" ", "")
    return bool(compact) and all(char in "-:" for char in compact)


def parse_plan_to_blocks(text: str) -> list[dict[str, object]]:
    text = text.replace("\r\n", "\n").strip()
    if not text:
        return []

    raw_blocks: list[list[str]] = []
    current: list[str] = []
    for line in text.splitlines():
        stripped = line.rstrip()
        if stripped.strip():
            current.append(stripped)
        elif current:
            raw_blocks.append(current)
            current = []
    if current:
        raw_blocks.append(current)

    parsed: list[dict[str, object]] = []
    for lines in raw_blocks:
        if not lines:
            continue

        if all(is_table_line(line) for line in lines):
            rows = [parse_table_row(line) for line in lines]
            rows = [row for row in rows if row and not is_markdown_table_separator(row)]
            if rows:
                parsed.append({"type": "table", "rows": rows})
            continue

        if all(is_list_line(line) for line in lines):
            items = [clean_markdown_text(line) for line in lines]
            parsed.append({"type": "list", "items": items})
            continue

        first_line = lines[0].strip()
        if is_markdown_heading_line(first_line) or is_numbered_heading_line(first_line) or (
            len(lines) > 1 and is_plain_heading_candidate(first_line) and not is_table_line(first_line) and not is_list_line(first_line)
        ):
            parsed.append({"type": "heading", "text": clean_markdown_text(first_line)})
            remainder = "\n".join(lines[1:]).strip()
            if remainder:
                parsed.extend(parse_plan_to_blocks(remainder))
            continue

        parsed.append({"type": "paragraph", "text": clean_markdown_text(" ".join(line.strip() for line in lines))})

    return parsed


def build_plan_latex(full_plan_text: str, context: dict[str, str] | None = None) -> str:
    title, subtitle, sections = parse_plan_blocks(full_plan_text)
    language = infer_plan_language(title, subtitle)
    subtitle = standardised_export_subtitle(language)
    title_line_one, title_line_two, _coordinator_label, cover_note = build_cover_page_text(language)
    school_field, coordinator_field, has_school_name = build_title_block_values(language, context)
    section_parts: list[str] = []
    for heading, paragraphs in sections:
        if heading:
            section_parts.append(f"\\section*{{{latex_escape(heading)}}}")
        for paragraph in paragraphs:
            if paragraph:
                section_parts.append(latex_escape(paragraph))
                section_parts.append("")
        if heading and needs_fill_lines(heading, language):
            for fill_line in export_fill_lines():
                section_parts.append(latex_escape(fill_line))
            section_parts.append("")

    body = "\n".join(section_parts).strip()
    escaped_title_line_one = latex_escape(title_line_one)
    escaped_title_line_two = latex_escape(title_line_two)
    escaped_subtitle = latex_escape(subtitle)
    escaped_cover_note = latex_escape(cover_note)
    escaped_school_field = latex_escape(school_field)
    escaped_coordinator_field = latex_escape(coordinator_field)
    school_line = (
        f"{{\\fontsize{{17pt}}{{21pt}}\\selectfont\\bfseries\\sffamily {escaped_school_field}\\par}}\n"
        if has_school_name
        else f"{{\\fontsize{{15pt}}{{19pt}}\\selectfont\\sffamily {escaped_school_field}\\par}}\n"
    )

    return (
        "\\documentclass[11pt]{article}\n"
        "\\usepackage[a4paper,margin=24mm,top=26mm,bottom=26mm]{geometry}\n"
        "\\usepackage{fontspec}\n"
        "\\usepackage{microtype}\n"
        "\\usepackage{titlesec}\n"
        "\\usepackage{setspace}\n"
        "\\usepackage{parskip}\n"
        "\\usepackage{ragged2e}\n"
        "\\usepackage{xcolor}\n"
        "\\usepackage[hidelinks]{hyperref}\n"
        "\\IfFontExistsTF{TeX Gyre Pagella}{\\setmainfont{TeX Gyre Pagella}}{\\IfFontExistsTF{Times New Roman}{\\setmainfont{Times New Roman}}{\\setmainfont{Latin Modern Roman}}}\n"
        "\\IfFontExistsTF{TeX Gyre Heros}{\\setsansfont{TeX Gyre Heros}}{\\IfFontExistsTF{Helvetica Neue}{\\setsansfont{Helvetica Neue}}{\\setsansfont{Latin Modern Sans}}}\n"
        "\\setstretch{1.16}\n"
        "\\setlength{\\parskip}{0.85em}\n"
        "\\setlength{\\parindent}{0pt}\n"
        "\\definecolor{TYHeading}{HTML}{203247}\n"
        "\\titleformat{\\section}{\\Large\\bfseries\\sffamily\\color{TYHeading}}{}{0pt}{}\n"
        "\\titlespacing*{\\section}{0pt}{1.8em}{0.55em}\n"
        "\\raggedbottom\n"
        "\\pagestyle{plain}\n"
        "\\begin{document}\n"
        "\\thispagestyle{empty}\n"
        "\\begin{center}\n"
        "\\vspace*{1.9cm}\n"
        f"{{\\fontsize{{22pt}}{{28pt}}\\selectfont\\bfseries\\sffamily {escaped_title_line_one}\\par}}\n"
        "\\vspace{0.35em}\n"
        f"{{\\fontsize{{22pt}}{{28pt}}\\selectfont\\bfseries\\sffamily {escaped_title_line_two}\\par}}\n"
        "\\vspace{1.15em}\n"
        f"{{\\large\\itshape {escaped_subtitle}\\par}}\n"
        "\\vspace{1.7em}\n"
        f"{school_line}"
        "\\vspace{0.9em}\n"
        f"{{\\normalsize\\sffamily {escaped_coordinator_field}\\par}}\n"
        "\\vspace{1.4em}\n"
        f"{{\\normalsize\\itshape {escaped_cover_note}\\par}}\n"
        "\\vfill\n"
        "\\rule{0.82\\textwidth}{0.6pt}\\par\n"
        "\\end{center}\n"
        "\\newpage\n"
        "\\RaggedRight\n"
        f"{body}\n"
        "\\end{document}\n"
    )


def resolve_reportlab_fonts():
    from pathlib import Path as _Path

    import reportlab

    fonts_dir = _Path(reportlab.__file__).resolve().parent / "fonts"
    font_sets = [
        (
            fonts_dir / "Vera.ttf",
            fonts_dir / "VeraBd.ttf",
            fonts_dir / "VeraIt.ttf",
            "TYVera",
            "TYVeraBold",
            "TYVeraItalic",
        ),
        (
            _Path("/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf"),
            _Path("/usr/share/fonts/truetype/dejavu/DejaVuSerif-Bold.ttf"),
            _Path("/usr/share/fonts/truetype/dejavu/DejaVuSerif-Italic.ttf"),
            "TYDejaVuSerif",
            "TYDejaVuSerifBold",
            "TYDejaVuSerifItalic",
        ),
        (
            _Path("/System/Library/Fonts/Supplemental/Times New Roman.ttf"),
            _Path("/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf"),
            _Path("/System/Library/Fonts/Supplemental/Times New Roman Italic.ttf"),
            "TYTimesNewRoman",
            "TYTimesNewRomanBold",
            "TYTimesNewRomanItalic",
        ),
        (
            _Path("/System/Library/Fonts/Supplemental/Georgia.ttf"),
            _Path("/System/Library/Fonts/Supplemental/Georgia Bold.ttf"),
            _Path("/System/Library/Fonts/Supplemental/Georgia Italic.ttf"),
            "TYGeorgia",
            "TYGeorgiaBold",
            "TYGeorgiaItalic",
        ),
    ]

    for regular_path, bold_path, italic_path, regular_name, bold_name, italic_name in font_sets:
        if regular_path.exists() and bold_path.exists() and italic_path.exists():
            return (regular_path, bold_path, italic_path, regular_name, bold_name, italic_name)
    return None


def validate_pdf_bytes(pdf_bytes: bytes) -> bytes:
    if not pdf_bytes.startswith(b"%PDF-"):
        raise RuntimeError("PDF export did not produce a valid PDF header.")
    if b"%%EOF" not in pdf_bytes[-2048:]:
        raise RuntimeError("PDF export did not produce a complete PDF trailer.")

    pdfinfo_path = shutil.which("pdfinfo")
    if pdfinfo_path:
        GENERATED_PLANS_DIR.mkdir(parents=True, exist_ok=True)
        with tempfile.NamedTemporaryFile(suffix=".pdf", dir=GENERATED_PLANS_DIR, delete=False) as handle:
            temp_pdf_path = Path(handle.name)
            handle.write(pdf_bytes)
        try:
            subprocess.run(
                [pdfinfo_path, str(temp_pdf_path)],
                check=True,
                capture_output=True,
                text=True,
            )
        finally:
            temp_pdf_path.unlink(missing_ok=True)
    return pdf_bytes


def validate_docx_bytes(docx_bytes: bytes) -> bytes:
    if not docx_bytes.startswith(b"PK"):
        raise RuntimeError("DOCX export did not produce a valid ZIP container.")
    with zipfile.ZipFile(BytesIO(docx_bytes)) as archive:
        names = set(archive.namelist())
        required = {"[Content_Types].xml", "word/document.xml"}
        if not required.issubset(names):
            raise RuntimeError("DOCX export is missing required document parts.")
    return docx_bytes


def build_pdf_fallback_bytes(full_plan_text: str, title: str, context: dict[str, str] | None = None) -> bytes:
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import (
            ListFlowable,
            ListItem,
            PageBreak,
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )
    except Exception as exc:
        raise RuntimeError("No PDF export backend is available locally.") from exc

    title_text, subtitle, _sections = parse_plan_blocks(full_plan_text)
    language = infer_plan_language(title_text, subtitle)
    subtitle = standardised_export_subtitle(language)
    _title_line_one, _title_line_two, _coordinator_label, cover_note = build_cover_page_text(language)
    school_line, _coordinator_line, has_school_name = build_title_block_values(language, context)
    context = normalise_template_context(context or {}, language)
    school_name = str(context.get("school_name", "")).strip()
    coordinator_name = str(context.get("ty_coordinator", "")).strip()
    clean_title = title_text or ("Plean Bliantúil na hIdirbhliana" if language == "ga" else "Transition Year Annual Plan")
    school_header = (
        school_line
        if school_name
        else ("Scoil: ______________________" if language == "ga" else "School: ______________________")
    )
    prepared_for_header = (
        f"Ullmhaithe do {coordinator_name}" if language == "ga" else f"Prepared for {coordinator_name}"
    ) if coordinator_name else ""

    body_text = full_plan_text
    title_block = f"{title_text}\n{subtitle}".strip()
    if title_block and body_text.startswith(title_block):
        body_text = body_text[len(title_block) :].lstrip()
    body_blocks = parse_plan_to_blocks(body_text)
    if prepared_for_header and body_blocks:
        first_block = body_blocks[0]
        if first_block.get("type") == "paragraph" and str(first_block.get("text", "")).strip() == prepared_for_header:
            body_blocks = body_blocks[1:]

    regular_font = "Helvetica"
    bold_font = "Helvetica-Bold"
    italic_font = "Helvetica-Oblique"
    font_set = resolve_reportlab_fonts()
    if font_set:
        regular_font_path, bold_font_path, italic_font_path, regular_font, bold_font, italic_font = font_set
        pdfmetrics.registerFont(TTFont(regular_font, str(regular_font_path)))
        pdfmetrics.registerFont(TTFont(bold_font, str(bold_font_path)))
        pdfmetrics.registerFont(TTFont(italic_font, str(italic_font_path)))

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=24 * mm,
        rightMargin=24 * mm,
        topMargin=24 * mm,
        bottomMargin=24 * mm,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TYTitle",
        parent=styles["Title"],
        fontName=bold_font,
        fontSize=21,
        leading=27,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#203247"),
        spaceAfter=8,
    )
    subtitle_style = ParagraphStyle(
        "TYSubtitle",
        parent=styles["Normal"],
        fontName=italic_font,
        fontSize=12,
        leading=16,
        alignment=TA_CENTER,
        spaceAfter=18,
    )
    cover_note_style = ParagraphStyle(
        "TYCoverNote",
        parent=styles["Normal"],
        fontName=italic_font,
        fontSize=11,
        leading=15,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#4A5560"),
        spaceAfter=10,
    )
    school_style = ParagraphStyle(
        "TYSchool",
        parent=styles["Normal"],
        fontName=bold_font if has_school_name else regular_font,
        fontSize=13 if has_school_name else 12,
        leading=18,
        alignment=TA_CENTER,
        spaceAfter=8,
    )
    coordinator_style = ParagraphStyle(
        "TYCoordinator",
        parent=styles["Normal"],
        fontName=italic_font,
        fontSize=11.5,
        leading=16,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#4A5560"),
        spaceAfter=14,
    )
    heading_style = ParagraphStyle(
        "TYHeading",
        parent=styles["Heading1"],
        fontName=bold_font,
        fontSize=15,
        leading=20,
        textColor=colors.HexColor("#203247"),
        spaceBefore=16,
        spaceAfter=8,
    )
    body_style = ParagraphStyle(
        "TYBody",
        parent=styles["BodyText"],
        fontName=regular_font,
        fontSize=11,
        leading=15,
        spaceAfter=9,
    )
    bullet_style = ParagraphStyle(
        "TYBullet",
        parent=body_style,
        leftIndent=12,
        firstLineIndent=0,
        spaceAfter=4,
    )
    table_cell_style = ParagraphStyle(
        "TYTableCell",
        parent=body_style,
        fontSize=10,
        leading=13,
        spaceAfter=0,
    )
    table_header_style = ParagraphStyle(
        "TYTableHeader",
        parent=table_cell_style,
        fontName=bold_font,
    )
    placeholder_style = ParagraphStyle(
        "TYPlaceholder",
        parent=body_style,
        fontName=italic_font,
        textColor=colors.HexColor("#68707A"),
    )

    story = [
        Spacer(1, 38),
        Paragraph(html.escape(clean_title), title_style),
        Paragraph(html.escape(subtitle), subtitle_style),
        Spacer(1, 10),
        Paragraph(html.escape(school_header), school_style),
    ]
    if prepared_for_header:
        story.append(Paragraph(html.escape(prepared_for_header), coordinator_style))
    story.extend([
        Paragraph(html.escape(cover_note), cover_note_style),
        PageBreak(),
    ])

    current_heading = ""
    for block in body_blocks + [{"type": "_end"}]:
        block_type = str(block.get("type", ""))
        if block_type in {"heading", "_end"} and current_heading and needs_fill_lines(current_heading, language):
            for fill_line in export_fill_lines():
                story.append(Paragraph(html.escape(fill_line), placeholder_style))
            current_heading = ""
        if block_type == "_end":
            continue
        if block_type == "heading":
            current_heading = str(block.get("text", "")).strip()
            if current_heading:
                story.append(Paragraph(html.escape(current_heading), heading_style))
        elif block_type == "paragraph":
            paragraph = str(block.get("text", "")).strip()
            if paragraph:
                style = placeholder_style if set(paragraph) <= {"_"} else body_style
                story.append(Paragraph(html.escape(paragraph).replace("\n", "<br/>"), style))
        elif block_type == "list":
            items = [str(item).strip() for item in block.get("items", []) if str(item).strip()]
            if items:
                list_flowable = ListFlowable(
                    [
                        ListItem(Paragraph(html.escape(item), bullet_style))
                        for item in items
                    ],
                    bulletType="bullet",
                    leftIndent=18,
                )
                story.append(list_flowable)
                story.append(Spacer(1, 6))
        elif block_type == "table":
            raw_rows = block.get("rows", [])
            rows = [[str(cell).strip() for cell in row] for row in raw_rows if isinstance(row, list)]
            if rows:
                max_cols = max(len(row) for row in rows)
                padded_rows = [row + [""] * (max_cols - len(row)) for row in rows]
                table_data = []
                for row_index, row in enumerate(padded_rows):
                    style = table_header_style if row_index == 0 else table_cell_style
                    table_data.append([Paragraph(html.escape(cell), style) for cell in row])
                table = Table(table_data, repeatRows=1, hAlign="LEFT")
                table.setStyle(
                    TableStyle(
                        [
                            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EDF2")),
                            ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#203247")),
                            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#B7C1CB")),
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                            ("LEFTPADDING", (0, 0), (-1, -1), 6),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                            ("TOPPADDING", (0, 0), (-1, -1), 5),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                        ]
                    )
                )
                story.append(table)
                story.append(Spacer(1, 8))

    doc.build(story)
    return validate_pdf_bytes(buffer.getvalue())


def build_docx_bytes(full_plan_text: str, context: dict[str, str] | None = None) -> bytes:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches
        from docx.shared import Pt
    except Exception as exc:  # pragma: no cover
        raise RuntimeError("python-docx is not available") from exc

    title, subtitle, sections = parse_plan_blocks(full_plan_text)
    language = infer_plan_language(title, subtitle)
    subtitle = standardised_export_subtitle(language)
    title_line_one, title_line_two, _coordinator_label, cover_note = build_cover_page_text(language)
    document = Document()

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Aptos"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.space_after = Pt(9)
    normal_style.paragraph_format.line_spacing = 1.16

    title_style = document.styles["Title"]
    title_style.font.name = "Aptos Display"
    title_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
    title_style.font.size = Pt(20)
    title_style.font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(5)

    heading_style = document.styles["Heading 1"]
    heading_style.font.name = "Aptos Display"
    heading_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
    heading_style.font.size = Pt(14)
    heading_style.font.bold = True
    heading_style.paragraph_format.space_before = Pt(18)
    heading_style.paragraph_format.space_after = Pt(6)

    for section in document.sections:
        section.top_margin = Inches(0.92)
        section.bottom_margin = Inches(0.92)
        section.left_margin = Inches(0.95)
        section.right_margin = Inches(0.95)

    for cover_title_line in (title_line_one, title_line_two):
        title_paragraph = document.add_paragraph(cover_title_line, style="Title")
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if subtitle:
        subtitle_paragraph = document.add_paragraph(subtitle, style="Normal")
        subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_paragraph.runs[0]
        subtitle_run.italic = True
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.name = "Aptos"
        subtitle_paragraph.paragraph_format.space_after = Pt(20)

    school_line, coordinator_line, has_school_name = build_title_block_values(language, context)
    school_paragraph = document.add_paragraph(school_line, style="Title" if has_school_name else "Normal")
    school_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    school_run = school_paragraph.runs[0]
    school_run.font.name = "Aptos Display" if has_school_name else "Aptos"
    school_run.font.size = Pt(17 if has_school_name else 13)
    school_paragraph.paragraph_format.space_after = Pt(14)

    coordinator_paragraph = document.add_paragraph(style="Normal")
    coordinator_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label, value = coordinator_line.split(":", 1)
    label_run = coordinator_paragraph.add_run(f"{label}:")
    label_run.bold = True
    value_run = coordinator_paragraph.add_run(value)
    value_run.font.name = "Aptos"
    coordinator_paragraph.paragraph_format.space_after = Pt(18)

    cover_note_paragraph = document.add_paragraph(cover_note, style="Normal")
    cover_note_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_note_run = cover_note_paragraph.runs[0]
    cover_note_run.italic = True
    cover_note_paragraph.paragraph_format.space_after = Pt(24)

    rule_paragraph = document.add_paragraph()
    rule_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAB2BD")
    p_borders.append(bottom)
    rule_paragraph._p.get_or_add_pPr().append(p_borders)
    rule_paragraph.paragraph_format.space_after = Pt(0)

    document.add_page_break()

    for heading, paragraphs in sections:
        if heading:
            heading_paragraph = document.add_heading(heading, level=1)
            heading_paragraph.paragraph_format.space_before = Pt(18)
            heading_paragraph.paragraph_format.space_after = Pt(6)
        for paragraph in paragraphs:
            if paragraph:
                body_paragraph = document.add_paragraph(paragraph, style="Normal")
                body_paragraph.paragraph_format.space_after = Pt(9)
                body_paragraph.paragraph_format.line_spacing = 1.16
        if heading and needs_fill_lines(heading, language):
            for fill_line in export_fill_lines():
                clean_fill_line = fill_line.replace("• ", "", 1)
                fill_paragraph = document.add_paragraph(clean_fill_line, style="List Bullet")
                fill_paragraph.paragraph_format.space_after = Pt(5)

    output = BytesIO()
    document.save(output)
    return validate_docx_bytes(output.getvalue())


def build_pdf_with_lualatex(full_plan_text: str, context: dict[str, str] | None = None) -> bytes:
    lualatex_path = shutil.which("lualatex")
    if not lualatex_path:
        raise RuntimeError("LuaLaTeX is not installed in the current environment.")

    latex_source = build_plan_latex(full_plan_text, context=context)
    with tempfile.TemporaryDirectory(prefix="ty_plan_pdf_", dir=GENERATED_PLANS_DIR) as temp_dir:
        temp_path = Path(temp_dir)
        tex_path = temp_path / "ty_plan_export.tex"
        pdf_path = temp_path / "ty_plan_export.pdf"
        tex_cache_dir = temp_path / ".texlive-cache"
        tex_cache_dir.mkdir(exist_ok=True)
        tex_path.write_text(latex_source, encoding="utf-8")
        env = dict(os.environ)
        env.setdefault("TEXMFCACHE", str(tex_cache_dir))
        env.setdefault("TEXMFVAR", str(tex_cache_dir))
        env.setdefault("HOME", str(temp_path))
        subprocess.run(
            [lualatex_path, "-interaction=nonstopmode", "-halt-on-error", tex_path.name],
            cwd=temp_path,
            check=True,
            capture_output=True,
            text=True,
            env=env,
        )
        return validate_pdf_bytes(pdf_path.read_bytes())


def build_pdf_bytes(full_plan_text: str, title: str, context: dict[str, str] | None = None) -> bytes:
    backend_errors: list[str] = []

    try:
        pdf_bytes = build_pdf_fallback_bytes(full_plan_text, title, context=context)
        print("[ty-plan-pdf] backend=reportlab status=ok", flush=True)
        return pdf_bytes
    except Exception as exc:
        backend_errors.append(f"ReportLab backend failed: {exc}")
        print(f"[ty-plan-pdf] backend=reportlab status=failed reason={exc.__class__.__name__}", flush=True)

    try:
        pdf_bytes = build_pdf_with_lualatex(full_plan_text, context=context)
        print("[ty-plan-pdf] backend=lualatex status=ok", flush=True)
        return pdf_bytes
    except Exception as exc:
        backend_errors.append(f"LuaLaTeX backend failed: {exc}")
        print(f"[ty-plan-pdf] backend=lualatex status=failed reason={exc.__class__.__name__}", flush=True)

    raise RuntimeError(" | ".join(backend_errors))


def main() -> None:
    st.set_page_config(page_title="Ciunas TY Planning Tool", page_icon="📘", layout="centered")

    st.title("Transition Year Planning Tool")
    st.write("Generate a structured TY annual plan in minutes - in English or Irish.")
    st.markdown(
        "- Create a full TY annual plan\n"
        "- Adapt it to your school\n"
        "- Download and edit immediately\n"
        "- Works in English and Irish"
    )
    st.info(
        "You can use this tool in English or Irish.\n\n"
        "Examples: `Create a TY plan`, `Create a TY plan in Irish`, `Cruthaigh plean Idirbhliana`."
    )
    render_source_layer_caption()

    mode = st.radio(
        "Choose a mode",
        ("Ask a TY Planning Question", "Generate a TY Annual Plan"),
        horizontal=True,
    )

    st.session_state.setdefault("generated_plan_result", None)
    st.session_state.setdefault("generated_plan_prompt", "Create a TY plan")
    st.session_state.setdefault("show_tailor_form", False)
    st.session_state.setdefault("download_format", "PDF")
    st.session_state.setdefault("download_unlocked", False)
    st.session_state.setdefault("lead_name", "")
    st.session_state.setdefault("lead_email", "")
    st.session_state.setdefault("planner_coordinator_name", "")

    if mode == "Ask a TY Planning Question":
        input_label = "Ask a TY planning question in English or Irish"
        input_placeholder = "For example: How should a school structure a Transition Year programme? / Conas ba chóir clár Idirbhliana a struchtúrú?"
        submit_label = "Ask"
        loading_message = "Finding grounded TY guidance..."
        empty_message = "Please enter a TY planning question before continuing."
        idle_message = "Enter a TY planning question in English or Irish to get grounded guidance."
    else:
        input_label = "Optional: add a focus or context"
        input_placeholder = "e.g. wellbeing focus, small rural school, student leadership, as Gaeilge"
        submit_label = "Generate TY Plan"
        loading_message = "Generating TY plan..."
        empty_message = ""
        idle_message = (
            "Generate a full TY Annual Plan in English or Irish, with optional context to shape the draft."
        )

    with st.form("ty_question_form", clear_on_submit=False):
        user_input = st.text_area(
            input_label,
            placeholder=input_placeholder,
            height=120,
            help=(
                "You can enter your own context or just generate a plan. You can ask in English or Irish. "
                "Examples: Create a TY plan, Create a TY plan focused on wellbeing, Cruthaigh plean Idirbhliana."
                if mode == "Generate a TY Annual Plan"
                else "You can ask in English or Irish."
            ),
        )
        name = ""
        if mode == "Generate a TY Annual Plan":
            name = st.text_input(
                "Name (optional)",
                value=st.session_state.get("planner_coordinator_name", ""),
                placeholder="e.g. Mary",
                help="If provided, the plan can be lightly personalised for the TY coordinator.",
            )
        submitted = st.form_submit_button(submit_label)

    if not submitted:
        if mode == "Ask a TY Planning Question" or not st.session_state.get("generated_plan_result"):
            st.info(idle_message)
            return
        result = st.session_state["generated_plan_result"]
        question = str(st.session_state.get("generated_plan_prompt", "Create a TY plan"))
    else:
        user_input = user_input.strip()
        if mode == "Ask a TY Planning Question":
            question = user_input
            if not question:
                st.warning(empty_message)
                return
        else:
            coordinator_name = name.strip()
            question = build_download_prompt(user_input, coordinator_name=coordinator_name)
            st.session_state["planner_coordinator_name"] = coordinator_name

        try:
            with st.spinner(loading_message):
                result = answer_question(
                    question,
                    coordinator_name=(
                        st.session_state.get("planner_coordinator_name", "").strip() or None
                        if mode == "Generate a TY Annual Plan"
                        else None
                    ),
                )
        except Exception as exc:  # pragma: no cover
            if mode == "Generate a TY Annual Plan":
                st.error("TY plan generation failed. Please try again.")
            else:
                st.error("The guidance engine hit a runtime error. Please try again.")
            return

        if mode == "Generate a TY Annual Plan":
            st.session_state["generated_plan_result"] = result
            st.session_state["generated_plan_prompt"] = question
            st.session_state["show_tailor_form"] = False

    answer = str(result.get("answer", "")).strip()
    key_points = result.get("key_points") or []
    sources = result.get("sources") or []
    evidence_note = str(result.get("evidence_note", "")).strip()
    answer_mode = str(result.get("answer_mode", "")).strip()
    generation_source = str(result.get("generation_source", "")).strip()
    model_used = str(result.get("model_used", "")).strip()
    full_plan_text = answer if mode == "Generate a TY Annual Plan" and is_generator_result(answer_mode) else ""

    if not answer and not key_points and not sources:
        if mode == "Generate a TY Annual Plan":
            st.warning("No TY plan could be generated from the current template mode.")
        else:
            st.warning("No grounded answer could be assembled for this question from the current source set.")
        return

    if mode == "Generate a TY Annual Plan":
        if not full_plan_text:
            st.error("Generator mode did not return a full TY plan, so export has been disabled for this response.")
            return
        unlocked = st.session_state.get("download_unlocked", False)
        if unlocked:
            render_generated_plan(full_plan_text)
            st.caption("You can download this plan or improve it further for your school.")
        else:
            render_plan_preview(full_plan_text)
            st.info("Continue to full plan by unlocking below.")
        output_language = infer_output_language(answer_mode, question)
        markdown_path, _text_path = save_generated_plan(full_plan_text, output_language)
        plan_title = full_plan_text.splitlines()[0].strip() if full_plan_text.splitlines() else "TY Annual Plan"
        prompt_context = parse_template_context(question)
        pdf_bytes: bytes | None = None
        pdf_error = ""
        try:
            pdf_bytes = build_pdf_bytes(full_plan_text, plan_title, context=prompt_context)
        except Exception as exc:  # pragma: no cover
            pdf_error = str(exc)
        docx_bytes: bytes | None = None
        docx_error = ""
        try:
            docx_bytes = build_docx_bytes(full_plan_text, context=prompt_context)
        except Exception as exc:  # pragma: no cover
            docx_error = str(exc)
        export_char_count = len(full_plan_text)
        print(
            f"[ty-plan-export] mode={answer_mode} language={output_language} chars={export_char_count}",
            flush=True,
        )
        st.markdown("### Download your formatted TY plan")
        st.caption("The downloadable version is formatted and ready to edit and share.")
        st.caption("Includes a clean layout suitable for school use.")
        st.caption(f"Export source length: {export_char_count} characters.")
        pdf_runtime_available, pdf_runtime_detail = describe_pdf_runtime_support()
        if SHOW_PDF_DIAGNOSTIC:
            st.caption(
                f"Temporary diagnostic: PDF runtime support is {'available' if pdf_runtime_available else 'unavailable'}."
            )
            if pdf_runtime_detail:
                st.caption(f"Temporary diagnostic: PDF runtime detail: {pdf_runtime_detail}")
            if pdf_error:
                st.caption(f"Temporary diagnostic: PDF unavailable reason: {pdf_error}")
        if SHOW_DOCX_DIAGNOSTIC:
            docx_runtime_available = bool(docx_bytes)
            st.caption(
                f"Temporary diagnostic: DOCX runtime support is {'available' if docx_runtime_available else 'unavailable'}."
            )
            if not docx_runtime_available:
                st.caption(
                    "Temporary diagnostic: DOCX unavailable reason: "
                    f"{docx_error or 'build_docx_bytes returned no content.'}"
                )
        if not unlocked:
            with st.form("download_unlock_form", clear_on_submit=False):
                lead_email = st.text_input(
                    "Email address",
                    value=st.session_state.get("lead_email", ""),
                    placeholder="e.g. teacher@school.ie",
                )
                lead_name = st.text_input(
                    "Name (optional)",
                    value=st.session_state.get("lead_name", ""),
                    placeholder="e.g. Mary",
                )
                st.caption("Enter your email to unlock the download and receive practical TY planning supports.")
                unlock_submitted = st.form_submit_button("Unlock download")

            if unlock_submitted:
                cleaned_email = lead_email.strip()
                cleaned_name = lead_name.strip()
                if not looks_like_email(cleaned_email):
                    st.warning("Please enter a valid email address to unlock the download.")
                else:
                    leads_path = save_lead_row(
                        {
                            "timestamp": datetime.now().isoformat(timespec="seconds"),
                            "email": cleaned_email,
                            "name": cleaned_name,
                            "mode": "Generate a TY Annual Plan",
                            "language": output_language,
                            "school_name": str(prompt_context.get("school_name", "")),
                            "school_type": str(prompt_context.get("school_type", "")),
                            "priorities": str(prompt_context.get("priorities", "")),
                        }
                    )
                    mailchimp_ok, mailchimp_detail = sync_mailchimp_lead(cleaned_email, cleaned_name or None)
                    if mailchimp_ok:
                        LOGGER.info("TY planner lead synced to Mailchimp for %s", cleaned_email)
                    else:
                        LOGGER.warning(
                            "TY planner Mailchimp sync failed for %s. Local CSV backup saved at %s. Detail: %s",
                            cleaned_email,
                            leads_path,
                            mailchimp_detail,
                        )
                    st.session_state["lead_email"] = cleaned_email
                    st.session_state["lead_name"] = cleaned_name
                    st.session_state["download_unlocked"] = True
                    st.rerun()
        else:
            if pdf_bytes and docx_bytes:
                download_options = ["PDF", "Word (.docx)"]
            elif pdf_bytes:
                download_options = ["PDF", "Markdown (.md fallback)"]
            elif docx_bytes:
                download_options = ["Word (.docx)", "Markdown (.md fallback)"]
            else:
                download_options = ["Markdown (.md fallback)"]
            current_format = st.session_state.get("download_format", "PDF")
            default_index = download_options.index(current_format) if current_format in download_options else 0
            download_format = st.selectbox("Choose format:", download_options, index=default_index, key="download_format")
            if download_format == "PDF" and pdf_bytes:
                pdf_path = save_export_file(pdf_bytes, output_language, "pdf", "pdf")
                st.caption(
                    "PDF export is generated directly from the full plan shown on screen."
                )
                st.download_button(
                    "Download plan",
                    data=pdf_bytes,
                    file_name=pdf_path.name,
                    mime="application/pdf",
                )
            elif download_format == "Word (.docx)" and docx_bytes:
                docx_path = save_export_file(docx_bytes or b"", output_language, "docx", "docx")
                st.caption(
                    "Word export preserves the title and section hierarchy for editing."
                )
                st.download_button(
                    "Download plan",
                    data=docx_bytes,
                    file_name=docx_path.name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            else:
                markdown_bytes = markdown_path.read_bytes()
                markdown_export_path = save_export_file(markdown_bytes, output_language, "markdown", "md")
                fallback_reasons = []
                if pdf_error:
                    fallback_reasons.append("PDF export is temporarily unavailable for this session.")
                if docx_error:
                    fallback_reasons.append("Word export is temporarily unavailable for this session.")
                if fallback_reasons:
                    for fallback_reason in fallback_reasons:
                        st.caption(fallback_reason)
                else:
                    st.caption(
                        "An editable Markdown fallback is available for this session."
                    )
                st.download_button(
                    "Download plan",
                    data=markdown_bytes,
                    file_name=markdown_export_path.name,
                    mime="text/markdown",
                )
        if st.button("Improve this plan for my school"):
            st.session_state["show_tailor_form"] = True

        if st.session_state.get("show_tailor_form"):
            st.markdown("### Tailor this plan for your school")
            st.write("Add a few details to improve the plan. Leave anything blank if not relevant.")
            with st.form("tailor_plan_form", clear_on_submit=False):
                school_name = st.text_input(
                    "School name (optional)",
                    placeholder="e.g. St. Mary's College",
                )
                cohort_size = st.text_input(
                    "Approximate TY cohort size (optional)",
                    placeholder="e.g. 68 students",
                )
                school_type = st.text_input(
                    "School type",
                    placeholder="e.g. small rural mixed school",
                    help=(
                        "Examples: Gaelcholáiste, English-medium school, mixed school, girls' school, "
                        "boys' school, large urban school, small rural school, DEIS school."
                    ),
                )
                school_ethos = st.text_input(
                    "School ethos or character (optional)",
                    placeholder="e.g. Catholic ethos, Gaelscoil, ETB school",
                )
                priorities = st.text_area(
                    "Main priorities for TY this year",
                    height=80,
                    placeholder="e.g. wellbeing and student leadership",
                    help=(
                        "Examples: wellbeing, student leadership, work experience, confidence, "
                        "balance, engagement"
                    ),
                )
                existing_modules = st.text_area(
                    "Existing modules or programmes",
                    height=80,
                    placeholder="e.g. mini-company and Gaisce",
                    help=(
                        "Examples: mini-company, Gaisce, outdoor education, musical, enterprise projects"
                    ),
                )
                work_experience = st.text_input(
                    "Work experience timing",
                    placeholder="e.g. two-week block in January",
                    help=(
                        "Examples: January block, spread across the year, two-week block, not yet decided"
                    ),
                )
                additional_context = st.text_area(
                    "Anything else we should include? (optional)",
                    height=100,
                    placeholder=(
                        "Examples: school ethos or mission, fixed events or trips, staffing or timetable "
                        "constraints, existing TY traditions, local community links, particular student "
                        "needs, assessment or portfolio expectations"
                    ),
                    help=(
                        "You can add any extra detail that would make the plan more realistic for your school."
                    ),
                )
                language = st.radio("Language for the final plan", ("English", "Irish"), horizontal=True)
                improve_submitted = st.form_submit_button("Generate improved plan")

            if improve_submitted:
                improved_prompt = build_tailored_plan_prompt(
                    school_name=school_name.strip(),
                    cohort_size=cohort_size.strip(),
                    coordinator_name=st.session_state.get("planner_coordinator_name", "").strip(),
                    school_type=school_type.strip(),
                    school_ethos=school_ethos.strip(),
                    priorities=priorities.strip(),
                    existing_modules=existing_modules.strip(),
                    work_experience=work_experience.strip(),
                    additional_context=additional_context.strip(),
                    language=language.strip(),
                )
                try:
                    with st.spinner("Generating improved TY plan..."):
                        improved_result = answer_question(
                            improved_prompt,
                            coordinator_name=st.session_state.get("planner_coordinator_name", "").strip() or None,
                        )
                except Exception as exc:  # pragma: no cover
                    st.error("Improved TY plan generation failed. Please try again.")
                    st.code(str(exc))
                    return

                improved_answer = str(improved_result.get("answer", "")).strip()
                if not improved_answer:
                    st.warning("No improved TY plan could be generated from the provided context.")
                    return

                st.session_state["generated_plan_result"] = improved_result
                st.session_state["generated_plan_prompt"] = improved_prompt
                st.session_state["show_tailor_form"] = False
                st.rerun()
    else:
        st.subheader("Answer")
        st.write(answer or "No grounded answer text was returned.")

    if mode == "Ask a TY Planning Question" and key_points:
        st.subheader("Key Points")
        for point in key_points:
            layer, body = split_layer_prefix(str(point))
            if layer:
                st.markdown(f"- `{layer}` {body}")
            else:
                st.markdown(f"- {body}")

    if mode == "Ask a TY Planning Question" and sources:
        st.subheader("Sources")
        for source in sources:
            title = str(source.get("title", "Untitled source"))
            layer = str(source.get("source_layer", "unknown"))
            chunk_id = str(source.get("chunk_id", ""))
            source_note = str(source.get("source_note", "")).strip()
            line = f"- `{layer}` {title}"
            if chunk_id:
                line += f" ({chunk_id})"
            if source_note:
                line += f" [{source_note}]"
            st.markdown(line)

    if evidence_note and mode == "Ask a TY Planning Question":
        st.subheader("Evidence Note")
        st.write(evidence_note)

if __name__ == "__main__":
    main()
